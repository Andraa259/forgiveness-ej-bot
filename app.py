import streamlit as st
from docx import Document
import requests
import io
import openpyxl
import gspread
from google.oauth2.service_account import Credentials
from streamlit_scroll_to_top import scroll_to_here
import time

# --- KREDENSIAL & CONFIG ---
TOKEN = st.secrets["TOKEN"]
ID_USER_WORD = st.secrets.get("CHAT_ID_1") 
ID_USER_FULL = st.secrets.get("CHAT_ID_2")
GSHEET_URL = st.secrets["GSHEET_URL"]

# --- INITIALIZING SESSION STATE ---
if 'step' not in st.session_state:
    st.session_state.step = 0
if 'scroll_to_top' not in st.session_state:
    st.session_state.scroll_to_top = False
if 'master_data' not in st.session_state:
    st.session_state.master_data = {}
if 'p_nama' not in st.session_state:
    st.session_state.p_nama = ""
if 'p_kerja' not in st.session_state:
    st.session_state.p_kerja = ""
if 'saran_global' not in st.session_state:
    st.session_state.saran_global = ""
if 'submitted' not in st.session_state:
    st.session_state.submitted = False
if 'confirmed' not in st.session_state: 
    st.session_state.confirmed = False

# --- LOGIKA SCROLL SISTEM ---
if st.session_state.scroll_to_top:
    scroll_to_here(0, key=f'scroll_step_{st.session_state.step}') 
    st.session_state.scroll_to_top = False

def move_step(step_num):
    st.session_state.step = step_num
    st.session_state.scroll_to_top = True

def sync_data(txt, key_type):
    # Efisiensi: Sinkronisasi instan agar visual feedback tidak delay
    st.session_state.master_data[txt][key_type] = st.session_state[f"{key_type}_{txt}"]

# --- FUNGSI INTEGRASI ---
def get_gsheet_client():
    scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scope)
    return gspread.authorize(creds)

def simpan_ke_gsheets():
    try:
        client = get_gsheet_client()
        ss = client.open_by_url(GSHEET_URL)
        kategori_map = {"KEJELASAN": "kj", "RELEVANSI": "rel", "KESESUAIAN": "kes"}
        for sheet_name, key_data in kategori_map.items():
            ws = ss.worksheet(sheet_name)
            col_c = ws.col_values(3)
            target_row = 4
            for r in range(4, 34):
                if r > len(col_c) or not col_c[r-1]:
                    target_row = r
                    break
                target_row = r + 1
            if target_row <= 33:
                ws.update_cell(target_row, 2, st.session_state.p_nama)
                skor_urut = []
                for aspek in ["Pemaafan Diri", "Pemaafan Orang Lain", "Pemaafan Situasi"]:
                    for _, items in data_aspek[aspek]:
                        for txt in items:
                            val = st.session_state.master_data.get(txt, {key_data: 0})[key_data]
                            skor_urut.append(val)
                cells = ws.range(target_row, 3, target_row, 3 + len(skor_urut) - 1)
                for i, score in enumerate(skor_urut): cells[i].value = score
                ws.update_cells(cells)
        return True
    except: return False

def proses_excel_cvi():
    try:
        client = get_gsheet_client()
        ss = client.open_by_url(GSHEET_URL)
        wb = openpyxl.load_workbook("CVI Aiken Zuyy.xlsx")
        for sheet_name in ["KEJELASAN", "RELEVANSI", "KESESUAIAN"]:
            ws_gs = ss.worksheet(sheet_name)
            all_vals = ws_gs.get_all_values()
            ws_xl = wb[sheet_name]
            for idx, row_data in enumerate(all_vals[3:]):
                target_row = 4 + idx
                if target_row > 33 or not row_data[2]: break
                ws_xl.cell(row=target_row, column=2, value=row_data[1])
                for col_idx, val in enumerate(row_data[2:38]):
                    try: ws_xl.cell(row=target_row, column=3 + col_idx, value=int(float(val)))
                    except: ws_xl.cell(row=target_row, column=3 + col_idx, value=0)
        buf = io.BytesIO(); wb.save(buf); buf.seek(0)
        return buf
    except: return None

def kirim_telegram_multi(word_buf, excel_buf, nama_panelis):
    url = f"https://api.telegram.org/bot{TOKEN}/sendDocument"
    targets = [
        {"id": ID_USER_WORD, "files": [("docx", word_buf)]},
        {"id": ID_USER_FULL, "files": [("docx", word_buf), ("xlsx", excel_buf)]}
    ]
    for target in targets:
        for f_type, f_buf in target["files"]:
            if f_buf is not None:
                fname = f"Form Validasi Expert Judgement Forgiveness_{nama_panelis}.docx" if f_type == "docx" else f"CVI_Aiken_Kumulatif_{nama_panelis}.xlsx"
                f_buf.seek(0)
                requests.post(url, data={'chat_id': target["id"], 'caption': f"✅ {f_type.upper()} Masuk: {nama_panelis}"}, files={'document': (fname, f_buf)})

# --- UI STYLING ---
st.set_page_config(page_title="Expert Judgement", layout="centered")
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    html, body, [class*="st-"] { font-family: 'Inter', sans-serif; }

    /* Paksa label dan teks di dalam container aitem jadi PUTIH agar kontras */
    div[data-testid="stVerticalBlock"] div[style*="background-color"] label, 
    div[data-testid="stVerticalBlock"] div[style*="background-color"] p, 
    div[data-testid="stVerticalBlock"] div[style*="background-color"] b,
    div[data-testid="stVerticalBlock"] div[style*="background-color"] span {
        color: #FFFFFF !important;
    }

    .intro-card { background: linear-gradient(135deg, #1E3A8A 0%, #3B82F6 100%); color: white !important; padding: 30px 40px; border-radius: 24px; text-align: center; margin-bottom: 30px; }
    .def-box { background-color: #F0F9FF !important; color: #075985 !important; padding: 18px; border-radius: 12px; border-left: 6px solid #0EA5E9; margin-bottom: 20px; line-height: 1.6; }
    .indicator-header { background-color: #1E3A8A; color: white !important; padding: 12px; border-radius: 10px 10px 0 0; font-weight: bold; text-align: center; margin-top: 15px; }
    
    .stButton>button { border-radius: 12px; height: 55px; font-weight: bold; width: 100%; background-color: #1E3A8A !important; color: white !important; }
    .thanks-card { text-align: center; padding: 40px; background-color: #F8FAFC !important; border-radius: 20px; border: 1px solid #E2E8F0; margin-top: 50px; }
    </style>
    """, unsafe_allow_html=True)

DEF_OP = "Pemaafan adalah kemampuan individual dalam membingkai ulang terhadap suatu kesalahan yang dialami/dirasakan sehingga mampu berhenti menyalahkan diri sendiri dan melepaskan pikiran negatif tentang diri sendiri, memahami kesalahan orang lain seiring berjalannya waktu serta berhenti berpikir buruk tentang orang yang pernah menyakiti, dan mampu berdamai dengan keadaan buruk dalam hidup serta melepaskan pikiran negatif terhadap peristiwa yang berada di luar kendali."

# --- DATA INDIKATOR ---
data_aspek = {
    "Pemaafan Diri": [
        ("Indikator 1: Kemampuan untuk berhenti menyalahkan diri sendiri", [
            "Seiring waktu, saya bisa memaklumi kesalahan pribadi yang pernah dilakukan. (Favorable)",
            "Ketika membuat kesalahan, saya fokus pada perbaikan daripada terus menerus menyalahkan diri sendiri. (Favorable)",
            "Saya memilih untuk berdamai dengan kekurangan diri sendiri. (Favorable)",
            "Sulit bagi saya untuk berhenti menyalahkan diri sendiri. (Unfavorable)",
            "Muncul perasaan benci ketika saya mengingat kesalahan diri sendiri. (Unfavorable)",
            "Saya terjebak dalam penyesalan atas kegagalan diri sendiri. (Unfavorable)"
        ]),
        ("Indikator 2: Kesediaan untuk melepaskan pikiran negatif tentang diri", [
            "Pikiran negatif tentang diri sendiri mulai memudar seiring waktu. (Favorable)",
            "Saya dapat memahami diri sendiri atas kesalahan yang telah saya lakukan. (Favorable)",
            "Saat  ingatan yang mengganggu tentang diri sendiri muncul, saya mampu melepaskannya. (Favorable)",
            "Sulit bagi saya untuk berhenti memikirkan hal-hal buruk yang pernah menimpa diri sendiri. (Unfavorable)",
            "Pikiran tentang kesalahan diri sendiri terus muncul walaupun sudah berusaha melupakannya. (Unfavorable)",
            "Saya sering susah berkonsentrasi karena teringat pada kesalahan diri sendiri yang telah lalu. (Unfavorable)"
        ])
    ],
    "Pemaafan Orang Lain": [
        ("Indikator 3: Kemampuan untuk memahami kesalahan orang lain", [
            "Saya dapat memaklumi bahwa setiap orang pasti pernah melakukan kekeliruan. (Favorable)",
            "Saya mencoba memahami alasan dibalik tindakan orang lain yang telah menyakiti saya. (Favorable)",
            "Saya menyadari bahwa ada alasan tertentu yang membuat orang lain sulit untuk bertindak benar. (Favorable)",
            "Memandang orang yang menyakiti saya sebagai pribadi yang memiliki karakter buruk. (Unfavorable)",
            "Saya tidak bisa menerima alasan apapun dari orang yang telah mengecewakan saya. (Unfavorable)",
            "Sangat sulit bagi saya untuk mengerti mengapa seseorang berbuat jahat kepada saya. (Unfavorable)"
        ]),
        ("Indikator 4: Berhenti berpikir buruk tentang orang yang pernah menyakiti", [
            "Pikiran buruk terhadap orang yang pernah menyakiti saya perlahan mulai menghilang. (Favorable)",
            "Saya merasa sudah tidak lagi menyimpan kebencian terhadap orang yang pernah menyakiti saya. (Favorable)",
            "Mudah bagi saya melepaskan rasa benci yang tertuju pada orang yang pernah berbuat salah. (Favorable)",
            "Saya terus membayangkan hal-hal negatif terjadi pada orang yang telah menyakiti saya. (Unfavorable)",
            "Sulit bagi saya untuk menghilangkan pandangan negatif terhadap orang yang pernah berbuat salah. (Unfavorable)",
            "Rasa kesal muncul kembali setiap kali saya mengingat perlakuan orang yang menyakiti saya. (Unfavorable)"
        ])
    ],
    "Pemaafan Situasi": [
        ("Indikator 5: Kemampuan untuk berdamai dengan keadaan buruk dalam hidup", [
            "Seiring berjalannya waktu, saya mulai bisa menerima kenyataan pahit yang terjadi dalam hidup dengan lapang dada. (Favorable)",
            "Saya sadar untuk tidak menyalahkan nasib atas kejadian buruk yang menimpa. (Favorable)",
            "Mampu menerima kenyataan bahwa hidup tidak selalu berjalan sesuai dengan rencana saya. (Favorable)",
            "Saya merasa semesta tidak adil karena terus memberikan cobaan yang berat. (Unfavorable)",
            "Sering merasa terjebak dalam nasib buruk yang seolah-olah tidak pernah berakhir di hidup saya. (Unfavorable)",
            "Terus-menerus mengeluhkan nasib buruk yang menimpa diri saya menjadi hal yang sulit untuk dihentikan. (Unfavorable)"
        ]),
        ("Indikator 6: Melepaskan pikiran negatif terhadap peristiwa luar kendali", [
            "Pikiran tentang kejadian buruk di masa lalu tidak lagi mengganggu saya untuk berkonsentrasi sehari-hari. (Favorable)",
            "Saya merasa sudah bisa berdamai dengan bayangan tentang masa-masa sulit yang pernah dialami. (Favorable)",
            "Saya mampu mengalihkan fokus dari peristiwa yang mengecewakan ke hal-hal yang lebih produktif. (Favorable)",
            "Sangat sulit bagi saya untuk tidak memikirkan kegagalan yang pernah dialami. (Unfavorable)",
            "Saya merasa terjebak dalam memori tentang kejadian buruk yang pernah saya alami. (Unfavorable)",
            "Bayangan mengenai ketidakadilan hidup di masa lalu sering kali muncul tanpa bisa saya kendalikan. (Unfavorable)"
        ])
    ]
}

# --- ALUR APLIKASI ---

# PROGRESS BAR (Muncul setelah Step 0)
if 0 < st.session_state.step < 5:
    st.progress((st.session_state.step) / 5.0)

# STEP 0: INTRO PORTAL AESTHETIC
if st.session_state.step == 0:
    st.markdown("""
        <div class='intro-card'>
            <h1 style='font-size: 3rem; margin-bottom: 10px;'>Expert Judgement Portal</h1>
            <p style='font-size: 1.2rem; opacity: 0.9;'>Instrument Validation for Forgiveness Scale Development</p>
            <div style='margin-top: 30px; font-weight: bold; letter-spacing: 3px; border: 2px solid white; display: inline-block; padding: 10px 20px; border-radius: 50px;'>WELCOME</div>
        </div>
    """, unsafe_allow_html=True)
    st.write("---")
    st.write("Selamat datang di sistem validasi instrumen. Partisipasi Anda sangat krusial dalam menjamin kualitas alat ukur penelitian ini.")
    if st.button("Mulai Sekarang ➔"): move_step(1); st.rerun()

# STEP 1: PETUNJUK & IDENTITAS
elif st.session_state.step == 1:
    st.title("⚖️ Form Validasi Expert Judgement")
    st.markdown(f"<div class='def-box'><b>Definisi Operasional:</b><br>{DEF_OP}</div>", unsafe_allow_html=True)
    st.subheader("📝 PETUNJUK PENGISIAN")
    st.info("Mohon dibaca sebelum memberikan penilaian")
    st.write("Sehubungan dengan upaya pengembangan instrumen penelitian mengenai tingkat pemaafan (forgiveness) pada mahasiswa, kami meminta Bapak/Ibu untuk menilai item-item yang telah kami susun, dari aspek :")
    st.markdown("""
    * **Kejelasan**: Kejelasan bahasa yang digunakan apakah sudah sesuai, jelas, dan mudah dipahami.
    * **Relevansi**: Relevansi aitem alat ukur yang disusun apakah sudah menggambarkan variabel.
    * **Kesesuaian**: Kesesuaian aitem yang disusun sudah sesuai dengan indikatornya.
    """)
    st.write("Penilaian dilakukan dengan memberikan angka 1-4. Skor **0** berarti Anda belum memberikan penilaian.")
    st.markdown("""
    0 = "Belum Diisi" | 1 = "Kurang" | 2 = "Cukup" | 3 = "Baik" | 4 = "Baik Sekali"
    """)
    st.write("---")
    st.session_state.p_nama = st.text_input("Nama Panelis", value=st.session_state.p_nama)
    st.session_state.p_kerja = st.text_input("Pekerjaan", value=st.session_state.p_kerja)
    
    if st.button("Lanjut ke Penilaian 🚀"):
        if not st.session_state.p_nama or not st.session_state.p_kerja:
            st.error("⚠️ Nama dan Pekerjaan wajib diisi!")
        else:
            move_step(2)
            st.rerun()

# STEP 2-4: PENILAIAN
elif st.session_state.step in [2, 3, 4]:
    aspek_list = {2: "Pemaafan Diri", 3: "Pemaafan Orang Lain", 4: "Pemaafan Situasi"}
    aspek_aktif = aspek_list[st.session_state.step]
    st.subheader(f"Aspek: {aspek_aktif}")

    for ind_name, items in data_aspek[aspek_aktif]:
        st.markdown(f"<div class='indicator-header'>{ind_name}</div>", unsafe_allow_html=True)
        for txt in items:
            if txt not in st.session_state.master_data:
                st.session_state.master_data[txt] = {"kj": 0, "rel": 0, "kes": 0, "ket": ""}
            
            d = st.session_state.master_data[txt]
            is_done = d["kj"] > 0 and d["rel"] > 0 and d["kes"] > 0
            
            if is_done:
                h_style = "background-color: #059669; border-left: 8px solid #047857; border: 1px solid #059669; border-bottom: none;"
                b_style = "background-color: #10B981; border-left: 8px solid #047857; border: 1px solid #10B981; border-top: none;"
                icon = "✅"
            else:
                h_style = "background-color: #334155; border-left: 8px solid #1E293B; border: 1px solid #334155; border-bottom: none;"
                b_style = "background-color: #475569; border-left: 8px solid #1E293B; border: 1px solid #475569; border-top: none;"
                icon = "⏳"

            st.markdown(f"<div style='{h_style} padding: 15px; border-radius: 8px 8px 0 0; margin-top: 25px;'><b style='color: white !important; font-size: 1.1rem;'>{icon} {txt}</b></div>", unsafe_allow_html=True)
            st.markdown(f"<div style='{b_style} padding: 10px 20px; border-radius: 0 0 8px 8px;'>", unsafe_allow_html=True)
            
            c1, c2, c3 = st.columns(3)
            with c1: st.session_state.master_data[txt]["kj"] = st.selectbox("Kejelasan", [0,1,2,3,4], index=d["kj"], key=f"kj_{txt}", on_change=sync_data, args=(txt, "kj"))
            with c2: st.session_state.master_data[txt]["rel"] = st.selectbox("Relevansi", [0,1,2,3,4], index=d["rel"], key=f"rel_{txt}", on_change=sync_data, args=(txt, "rel"))
            with c3: st.session_state.master_data[txt]["kes"] = st.selectbox("Kesesuaian", [0,1,2,3,4], index=d["kes"], key=f"kes_{txt}", on_change=sync_data, args=(txt, "kes"))
            st.session_state.master_data[txt]["ket"] = st.text_input("Keterangan per Aitem:", value=d["ket"], key=f"ket_{txt}")
            st.markdown("</div>", unsafe_allow_html=True)

    items_pg = [t for _, items in data_aspek[aspek_aktif] for t in items]
    errors = [t for t in items_pg if any(st.session_state.master_data[t][k] == 0 for k in ["kj", "rel", "kes"])]

    if st.session_state.step == 4:
        st.session_state.saran_global = st.text_area("Catatan/Saran Keseluruhan:", value=st.session_state.saran_global)

    nav1, nav2 = st.columns(2)
    with nav1:
        if st.button("⬅️ Kembali"): move_step(st.session_state.step - 1); st.rerun()
    with nav2:
        btn_label = "Lanjut ➡️" if st.session_state.step < 4 else "🚀 LANJUT KE PENGIRIMAN"
        if st.button(btn_label):
            if errors: st.error(f"⚠️ Ada {len(errors)} soal yang belum lengkap pada halaman ini.")
            else: move_step(5 if st.session_state.step == 4 else st.session_state.step + 1); st.rerun()

# STEP 5: KONFIRMASI
elif st.session_state.step == 5:
    st.title("Konfirmasi & Pengiriman")
    st.warning("Mohon periksa kembali penilaian Anda. Setelah dikirim, data tidak dapat diubah.")
    st.session_state.confirmed = st.checkbox("Saya menyatakan bahwa data yang saya masukkan sudah benar.")
    
    nav1, nav2 = st.columns(2)
    with nav1:
        if st.button("⬅️ Kembali ke Penilaian"): move_step(4); st.rerun()
    with nav2:
        if st.button("🚀 YA, KIRIM SEKARANG", disabled=not st.session_state.confirmed):
            with st.spinner("Sedang memproses..."):
                try:
                    simpan_ke_gsheets(); excel_buf = proses_excel_cvi()
                    doc = Document("Form Validasi Expert Judgement Ayinn Ver. 3.docx")
                    for p in doc.paragraphs:
                        if "Nama\t\t:" in p.text: p.text = f"Nama\t\t: {st.session_state.p_nama}"
                        if "Pekerjaan\t:" in p.text: p.text = f"Pekerjaan\t: {st.session_state.p_kerja}"
                    table = doc.tables[0]
                    for row in table.rows:
                        a_word = "".join(row.cells[2].text.split()).lower()
                        for t_ori, d in st.session_state.master_data.items():
                            if "".join(t_ori.split()).lower()[:60] in a_word:
                                row.cells[3].text, row.cells[4].text, row.cells[5].text, row.cells[6].text = str(d["kj"]), str(d["rel"]), str(d["kes"]), d["ket"]
                    if st.session_state.saran_global:
                        for row in table.rows:
                            if "Catatan" in row.cells[2].text: row.cells[2].text += "\n" + st.session_state.saran_global
                    word_buf = io.BytesIO(); doc.save(word_buf); word_buf.seek(0)
                    kirim_telegram_multi(word_buf, excel_buf, st.session_state.p_nama)
                    st.session_state.submitted = True; move_step(6); st.rerun()
                except Exception as e: st.error(f"Error: {e}")

# STEP 6: SELESAI
elif st.session_state.step == 6:
    st.balloons()
    st.markdown("""
        <div class='thanks-card'>
            <h1 style='color: #1E3A8A;'>Terima Kasih! ✨</h1>
            <p style='font-size: 1.2rem; color: #475569;'>
                Data penilaian Anda telah berhasil kami terima dan dikirimkan ke peneliti. 
                Kontribusi Anda sangat berharga bagi pengembangan instrumen penelitian ini.
            </p>
            <hr>
            <p style='font-style: italic; color: #64748b;'>Halaman ini dapat Anda tutup sekarang.</p>
        </div>
    """, unsafe_allow_html=True)
